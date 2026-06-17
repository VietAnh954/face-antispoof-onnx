import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_cell_background(cell, fill_color):
    """Sets background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets padding for table cells in twentieths of a point (dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_paragraph_with_spacing(doc, text="", style=None, space_before=Pt(0), space_after=Pt(6), line_spacing=1.15):
    """Creates a paragraph with controlled spacing and line height."""
    p = doc.add_paragraph(text, style=style)
    p_format = p.paragraph_format
    p_format.space_before = space_before
    p_format.space_after = space_after
    p_format.line_spacing = line_spacing
    return p

def main():
    doc = Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Off-black

    # Color Palette
    PRIMARY_COLOR = RGBColor(0x1F, 0x4E, 0x79) # Deep Navy
    SECONDARY_COLOR = RGBColor(0x5B, 0x9B, 0xD5) # Steel Blue
    ACCENT_COLOR = RGBColor(0x70, 0x30, 0xA0) # Dark Purple
    LIGHT_BG = "F2F5F8" # Hex string for table headers

    # 1. Document Title
    p_title = add_paragraph_with_spacing(doc, space_after=Pt(18))
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("HƯỚNG DẪN SỬ DỤNG VÀ KIỂM THỬ HỆ THỐNG\nPHÁT HIỆN GIẢ MẠO KHUÔN MẶT (FACE ANTI-SPOOFING)")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    # Author/Subtitle Info
    p_sub = add_paragraph_with_spacing(doc, space_after=Pt(24))
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tài liệu hướng dẫn triển khai, vận hành và kiểm thử hiệu năng tối ưu phần cứng (CPU, GPU, TensorRT)")
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True

    # 2. Section: Tổng quan hệ thống
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    r1 = h1.add_run("1. Tổng Quan Về Hệ Thống")
    r1.font.color.rgb = PRIMARY_COLOR
    r1.font.bold = True

    p = add_paragraph_with_spacing(doc)
    p.add_run("Hệ thống phát hiện giả mạo khuôn mặt (Face Anti-Spoofing - FAS) này được thiết kế để phân biệt giữa khuôn mặt thật của người đang đứng trước camera (Bona Fide Presentation) và khuôn mặt giả mạo (Presentation Attack) được trình chiếu qua màn hình điện thoại, máy tính bảng hoặc ảnh in màu.")

    p = add_paragraph_with_spacing(doc)
    p.add_run("Mô hình cốt lõi dựa trên kiến trúc mạng siêu nhẹ ").font.bold = False
    r_bold = p.add_run("MiniFASNetV2-SE")
    r_bold.font.bold = True
    p.add_run(" kết hợp với nhánh sinh ảnh tần số phổ Fourier (").font.bold = False
    p.add_run("Fourier Transform - FT").font.bold = True
    p.add_run(") trong quá trình huấn luyện nhằm phân tích cấu trúc tần số hạt nhiễu (moire patterns) trên màn hình hoặc độ tương phản đặc trưng của ảnh in.")

    # 3. Section: Cài đặt môi trường
    h2 = doc.add_heading(level=1)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    r2 = h2.add_run("2. Cài Đặt Môi Trường & Thư Viện")
    r2.font.color.rgb = PRIMARY_COLOR
    r2.font.bold = True

    p = add_paragraph_with_spacing(doc)
    p.add_run("Hệ thống hỗ trợ chạy trên cả Windows và Linux. Để đạt hiệu năng cao nhất, khuyến nghị chạy trên máy tính có GPU NVIDIA hỗ trợ CUDA.")

    p = add_paragraph_with_spacing(doc)
    p.add_run("Các bước cài đặt chi tiết trên Windows:").font.bold = True

    # List items for installation
    steps = [
        ("Cài đặt Python:", " Khuyến nghị sử dụng Python phiên bản 3.10 đến 3.13."),
        ("Cài đặt PyTorch với CUDA:", " Mở Windows PowerShell và chạy lệnh cài đặt PyTorch hỗ trợ GPU CUDA (tùy thuộc vào phiên bản driver của bạn):\nLệnh: `pip install torch torchvision --index-url https://download.pytorch.org/whl/cu124`"),
        ("Cài đặt ONNX và ONNX Runtime GPU:", " Dùng để chạy suy luận mô hình tối ưu dạng ONNX.\nLệnh: `pip install onnx onnxsim onnxruntime-gpu`"),
        ("Cài đặt TensorRT (tùy chọn tối ưu phần cứng):", " Cần cài đặt thư viện TensorRT của NVIDIA và Python API đi kèm (`pip install tensorrt`). Bản TensorRT khuyến nghị là 10.x hoặc 11.x."),
        ("Cài đặt các thư viện bổ trợ khác:", " Lệnh: `pip install numpy pandas scikit-learn opencv-python Pillow tqdm py-cpuinfo GPUtil psutil`"),
    ]

    for title, desc in steps:
        p_list = add_paragraph_with_spacing(doc, space_after=Pt(4), style='List Bullet')
        r_title = p_list.add_run(title)
        r_title.font.bold = True
        p_list.add_run(desc)

    # Note about Windows DLL
    p_note = add_paragraph_with_spacing(doc, space_before=Pt(6), space_after=Pt(6))
    r_note = p_note.add_run("Lưu ý xử lý lỗi trên Windows: ")
    r_note.font.bold = True
    r_note.font.color.rgb = ACCENT_COLOR
    p_note.add_run("Khi chạy ONNX Runtime GPU trên Windows, hệ thống có thể báo lỗi thiếu file DLL (ví dụ: cublasLt64_12.dll). Để khắc phục điều này, mã nguồn hệ thống tự động import thư viện `onnxruntime` trước tiên và thực hiện hàm `onnxruntime.preload_dlls(cuda=True, cudnn=True)` để tự động tải các DLL cần thiết từ thư viện CUDA đi kèm của Python.")

    # 4. Section: Chuẩn bị mô hình
    h3 = doc.add_heading(level=1)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    r3 = h3.add_run("3. Quy Trình Xuất Và Biên Dịch Mô Hình")
    r3.font.color.rgb = PRIMARY_COLOR
    r3.font.bold = True

    p = add_paragraph_with_spacing(doc)
    p.add_run("Hệ thống đi kèm các checkpoint huấn luyện định dạng PyTorch (.pth). Để chạy demo và kiểm thử hiệu năng, ta cần xuất chúng sang định dạng ONNX và biên dịch sang TensorRT Engine.")

    # Table of commands
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Định Dạng Mô Hình'
    hdr_cells[1].text = 'Mục Tiêu & Đặc Điểm'
    hdr_cells[2].text = 'Lệnh Chạy Thực Thi'

    for cell in hdr_cells:
        set_cell_background(cell, LIGHT_BG)
        set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = PRIMARY_COLOR

    commands_data = [
        (
            "ONNX Baseline (FP32)",
            "Xuất mô hình gốc PyTorch (.pth) sang định dạng ONNX tiêu chuẩn để chạy trên CPU/GPU.",
            "python scripts/export_onnx.py models/best/98.20/best_model.pth"
        ),
        (
            "ONNX Quantized (INT8)",
            "Lượng tử hóa động (dynamic quantization) giảm kích thước file mô hình xuống ~4 lần (~600KB) chạy trên CPU.",
            "python scripts/quantize_onnx.py models/best/98.20/best_model.pth"
        ),
        (
            "TensorRT Engine (FP16)",
            "Biên dịch sang dạng máy cơ sở (native engine) của NVIDIA, tối ưu hóa FP16 cho GPU Laptop/Card đồ họa.",
            "python scripts/export_tensorrt.py --model models/best/98.20/best_model.pth --output models/best_model_fp16.engine --fp16"
        ),
        (
            "TensorRT Engine (INT8)",
            "Biên dịch tối ưu hóa lượng tử hóa tĩnh tĩnh INT8 sử dụng QDQ nodes và bộ dữ liệu hiệu chuẩn (calibration data).",
            "python scripts/export_tensorrt.py --model models/best/98.20/best_model.pth --output models/best_model_int8.engine --int8 --calib_dir data_calib"
        )
    ]

    for fmt, desc, cmd in commands_data:
        row_cells = table.add_row().cells
        row_cells[0].text = fmt
        row_cells[1].text = desc
        row_cells[2].text = cmd
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            # Make commands monospace and slightly smaller
            if idx == 2:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9.5)
            # Make format titles bold
            elif idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    # 5. Section: Chạy Demo thời gian thực
    h4 = doc.add_heading(level=1)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    r4 = h4.add_run("4. Hướng Dẫn Chạy Demo Thời Gian Thực")
    r4.font.color.rgb = PRIMARY_COLOR
    r4.font.bold = True

    p = add_paragraph_with_spacing(doc)
    p.add_run("File `demo.py` cung cấp giao diện hiển thị hình ảnh trực tiếp từ webcam hoặc xử lý ảnh tĩnh từ file ngoài để phân loại.")

    p = add_paragraph_with_spacing(doc)
    p.add_run("Cách chạy demo trên webcam (mặc định):").font.bold = True
    p_cmd1 = add_paragraph_with_spacing(doc, space_after=Pt(4))
    r_cmd1 = p_cmd1.add_run("Lệnh: python demo.py")
    r_cmd1.font.name = 'Courier New'
    r_cmd1.font.bold = True

    p = add_paragraph_with_spacing(doc)
    p.add_run("Cách chạy demo trên một ảnh tĩnh cụ thể:").font.bold = True
    p_cmd2 = add_paragraph_with_spacing(doc, space_after=Pt(4))
    r_cmd2 = p_cmd2.add_run("Lệnh: python demo.py --image duong_dan_anh.jpg")
    r_cmd2.font.name = 'Courier New'
    r_cmd2.font.bold = True

    p = add_paragraph_with_spacing(doc)
    p.add_run("Các tham số dòng lệnh quan trọng:").font.bold = True

    params = [
        ("--camera:", " Chỉ số camera kết nối với PC, mặc định là 0."),
        ("--threshold:", " Ngưỡng xác suất phân loại (từ 0 đến 1), mặc định là 0.5. Tăng ngưỡng để thắt chặt bảo mật liveness."),
        ("--bbox_expansion_factor:", " Hệ số mở rộng bounding box của mặt trước khi đưa vào mô hình phân loại, mặc định là 1.5."),
        ("--liveness_model:", " Đường dẫn tới mô hình liveness .onnx hoặc .engine muốn dùng chạy thực tế."),
    ]
    for param_name, param_desc in params:
        p_list = add_paragraph_with_spacing(doc, space_after=Pt(4), style='List Bullet')
        r_param = p_list.add_run(param_name)
        r_param.font.bold = True
        p_list.add_run(param_desc)

    p = add_paragraph_with_spacing(doc)
    p.add_run("Phím tắt điều khiển khi giao diện demo mở ra:").font.bold = True
    p_keys = add_paragraph_with_spacing(doc, style='List Bullet')
    p_keys.add_run("Nhấn phím ").font.bold = False
    p_keys.add_run("q").font.bold = True
    p_keys.add_run(" để thoát chương trình.")
    
    p_keys2 = add_paragraph_with_spacing(doc, style='List Bullet')
    p_keys2.add_run("Nhấn phím ").font.bold = False
    p_keys2.add_run("i").font.bold = True
    p_keys2.add_run(" để bật/tắt bảng thông tin overlay hiển thị FPS, CPU, GPU và nhà cung cấp phần cứng (Execution Provider).")

    # 6. Section: Kiểm thử & Đánh giá
    h5 = doc.add_heading(level=1)
    h5.paragraph_format.space_before = Pt(12)
    h5.paragraph_format.space_after = Pt(6)
    r5 = h5.add_run("5. Hướng Dẫn Kiểm Thử Và Đánh Giá Hệ Thống")
    r5.font.color.rgb = PRIMARY_COLOR
    r5.font.bold = True

    p = add_paragraph_with_spacing(doc)
    p.add_run("Giảng viên thường yêu cầu số liệu chính xác để chấm điểm. Hệ thống cung cấp hai công cụ kiểm thử quan trọng:")

    # Sub-heading 5.1
    h5_1 = doc.add_heading(level=2)
    h5_1.paragraph_format.space_before = Pt(6)
    h5_1.paragraph_format.space_after = Pt(4)
    r5_1 = h5_1.add_run("5.1. Kiểm thử hiệu năng phần cứng (Hardware Benchmarking)")
    r5_1.font.color.rgb = SECONDARY_COLOR
    r5_1.font.bold = True

    p = add_paragraph_with_spacing(doc)
    p.add_run("Script `benchmark_hardware.py` giúp đo đạc chính xác thời gian suy luận (Latency) và số khung hình xử lý trên giây (FPS), đồng thời giám sát tải sử dụng CPU, RAM, GPU và VRAM.")

    p_cmd3 = add_paragraph_with_spacing(doc, space_after=Pt(4))
    r_cmd3 = p_cmd3.add_run("Lệnh chạy: python benchmark_hardware.py --iters 100")
    r_cmd3.font.name = 'Courier New'
    r_cmd3.font.bold = True

    p = add_paragraph_with_spacing(doc)
    p.add_run("Script này sẽ lần lượt kiểm thử các cấu hình: ONNX CPU, ONNX GPU, TensorRT FP16 và TensorRT INT8, sau đó xuất ra bảng so sánh chi tiết. Giúp chứng minh hiệu năng của GPU Preprocessing kết hợp CUDA Graphs.")

    # Sub-heading 5.2
    h5_2 = doc.add_heading(level=2)
    h5_2.paragraph_format.space_before = Pt(6)
    h5_2.paragraph_format.space_after = Pt(4)
    r5_2 = h5_2.add_run("5.2. Đánh giá độ chính xác theo chuẩn quốc tế ISO/IEC 30107-3")
    r5_2.font.color.rgb = SECONDARY_COLOR
    r5_2.font.bold = True

    p = add_paragraph_with_spacing(doc)
    p.add_run("Để báo cáo học thuật mang tính khoa học, ta cần đánh giá chất lượng phân loại liveness bằng các chỉ số tiêu chuẩn. Quy trình thực hiện gồm hai bước:")

    # Nested steps
    p_step1 = add_paragraph_with_spacing(doc, style='List Bullet')
    p_step1.add_run("Bước 1: Chuẩn bị dữ liệu kiểm thử khuôn mặt từ tập dữ liệu gốc CelebA-Spoof.").font.bold = True
    p_step1.add_run("\nChạy lệnh: `python scripts/prepare_calib_data.py` để trích xuất 250 ảnh khuôn mặt đã được cắt và căn chỉnh (125 ảnh thật, 125 ảnh giả mạo) lưu trữ trong thư mục `data_calib/`.")
    
    p_step2 = add_paragraph_with_spacing(doc, style='List Bullet')
    p_step2.add_run("Bước 2: Thực hiện đánh giá chéo giữa các mô hình.").font.bold = True
    p_step2.add_run("\nChạy lệnh: `python scripts/evaluate_metrics.py`. Script sẽ tải 250 ảnh từ thư mục `data_calib/`, đẩy qua các mô hình và tính toán các tỉ lệ sai số theo tiêu chuẩn ISO/IEC 30107-3.")

    # Explanation of metrics
    p = add_paragraph_with_spacing(doc)
    p.add_run("Các chỉ số học thuật được báo cáo bao gồm:").font.bold = True

    metrics_list = [
        ("Accuracy (Độ chính xác toàn cục):", " Phần trăm các mẫu ảnh được phân loại đúng trong tổng số ảnh."),
        ("APCER (Attack Presentation Classification Error Rate):", " Tỉ lệ lỗi phân loại giả mạo. APCER là tỉ lệ ảnh giả mạo bị mô hình nhận nhầm là người thật (càng thấp càng tốt cho bảo mật)."),
        ("BPCER (Bona Fide Presentation Classification Error Rate):", " Tỉ lệ lỗi phân loại người thật. BPCER là tỉ lệ ảnh người thật bị mô hình nhận nhầm là giả mạo (càng thấp càng tốt cho trải nghiệm người dùng)."),
        ("ACER (Average Classification Error Rate):", " Tỉ lệ lỗi phân loại trung bình. Được tính bằng trung bình cộng của APCER và BPCER: ACER = (APCER + BPCER) / 2. ACER càng nhỏ mô hình hoạt động càng chính xác."),
    ]
    for m_name, m_desc in metrics_list:
        p_list = add_paragraph_with_spacing(doc, space_after=Pt(4), style='List Bullet')
        r_m = p_list.add_run(m_name)
        r_m.font.bold = True
        p_list.add_run(m_desc)

    # Save Document
    output_path = "Huong_Dan_Su_Dung_He_Thong.docx"
    doc.save(output_path)
    print(f"[OK] Word Document successfully generated and saved to: {output_path}")

if __name__ == "__main__":
    main()
