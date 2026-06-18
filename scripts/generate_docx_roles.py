import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Sets background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets padding for table cells in twentieths of a point (dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_paragraph_with_spacing(doc, text="", style=None, space_before=Pt(0), space_after=Pt(6), line_spacing=1.15):
    """Creates a paragraph with controlled spacing and line height."""
    p = doc.add_paragraph(text, style=style)
    p_format = p.paragraph_format
    p_format.space_before = space_before
    p_format.space_after = space_after
    p_format.line_spacing = line_spacing
    return p

def create_member_1_doc():
    doc = Document()
    sections = doc.sections
    for s in sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1.0)

    # Styles
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    PRIMARY_COLOR = RGBColor(0x1F, 0x4E, 0x79) # Navy
    SECONDARY_COLOR = RGBColor(0x5B, 0x9B, 0xD5)

    # Header
    p_title = add_paragraph_with_spacing(doc, space_after=Pt(12))
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("BẢN PHÂN CÔNG NHIỆM VỤ THÀNH VIÊN 1: KỸ SƯ TRÍ TUỆ NHÂN TẠO (AI/DEEP LEARNING ENGINEER)")
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    p_role = add_paragraph_with_spacing(doc, space_after=Pt(18))
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role.add_run("Mức độ đóng góp: ").font.bold = True
    p_role.add_run("CHÍNH (LỚN NHẤT) - 40% khối lượng dự án").font.color.rgb = PRIMARY_COLOR

    # Section 1: Overview
    h = doc.add_heading(level=1)
    h.add_run("1. Mục Tiêu Chính").font.color.rgb = PRIMARY_COLOR
    p = add_paragraph_with_spacing(doc)
    p.add_run("Thành viên này chịu trách nhiệm xây dựng 'bộ não' AI cốt lõi cho hệ thống. Nhiệm vụ chính xoay quanh thiết kế mô hình học máy, chuẩn bị dữ liệu ảnh khuôn mặt thật/giả, lập trình quy trình huấn luyện (training) để AI biết cách phân biệt, và tối ưu hóa các hàm toán học đo lường sai số.")

    # Section 2: Files
    h = doc.add_heading(level=1)
    h.add_run("2. Các Tệp Tin Đảm Nhận Trong Dự Án").font.color.rgb = PRIMARY_COLOR
    
    files_data = [
        ("src/minifasv2/model.py", "Chứa định nghĩa cấu trúc toán học của mạng nơ-ron MiniFASNetV2-SE và khối Squeeze-and-Excitation (SE)."),
        ("src/minifasv2/data.py", "Quản lý việc đọc ảnh, thực hiện biến đổi Fourier tần số để hỗ trợ huấn luyện và bộ tạo dữ liệu đầu vào."),
        ("scripts/train.py", "Script khởi chạy quá trình huấn luyện mô hình trên GPU/CPU, cài đặt các siêu tham số học tập."),
        ("src/minifasv2/main.py", "Chứa toàn bộ vòng lặp huấn luyện, đánh giá sai số (loss) sau mỗi lượt (epoch), và lưu trữ trọng số."),
        ("models/best/98.20/best_model.pth", "Tệp lưu trữ toàn bộ các tham số trọng số AI đã được huấn luyện thành công.")
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Đường Dẫn Tệp Tin"
    hdr_cells[1].text = "Vai Trò Chi Tiết Trong Mã Nguồn"
    for cell in hdr_cells:
        set_cell_background(cell, "F2F5F8")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = PRIMARY_COLOR

    for path, desc in files_data:
        row_cells = table.add_row().cells
        row_cells[0].text = path
        row_cells[1].text = desc
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            if idx == 0:
                cell.paragraphs[0].runs[0].font.name = 'Courier New'
                cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                cell.paragraphs[0].runs[0].font.bold = True

    # Section 3: Tasks
    h = doc.add_heading(level=1)
    h.add_run("3. Nhiệm Vụ Chi Tiết Từng Bước").font.color.rgb = PRIMARY_COLOR
    
    tasks = [
        ("Nghiên cứu kiến trúc mạng MiniFASNetV2-SE:", " Thiết lập mạng nơ-ron tích chập (CNN) siêu nhẹ. Thiết kế các khối Depthwise Separable Convolutions giúp giảm số lượng tham số nhưng vẫn giữ được khả năng nhận diện hình ảnh."),
        ("Xử lý dữ liệu CelebA-Spoof:", " Đọc dữ liệu từ tập dữ liệu lớn CelebA-Spoof. Phân chia ảnh thành 2 lớp: lớp 0 (Ảnh mặt người thật) và lớp 1 (Ảnh mặt giả mạo qua điện thoại, ảnh in)."),
        ("Tích hợp phép Biến đổi Fourier (Fourier Transform):", " Viết hàm chuyển đổi ảnh mặt sang phổ tần số (Gray-scale Fourier spectrum) để AI phát hiện các tần số Moire của màn hình điện thoại hoặc hạt mực in."),
        ("Lập trình vòng lặp huấn luyện (Training Loop):", " Sử dụng hàm tối ưu Adam/SGD, thiết lập lịch giảm tốc độ học (learning rate scheduler) và lưu lại mô hình tốt nhất (.pth) khi độ chính xác trên tập kiểm thử đạt cao nhất.")
    ]
    for title, desc in tasks:
        p_list = add_paragraph_with_spacing(doc, style='List Bullet')
        r_title = p_list.add_run(title)
        r_title.font.bold = True
        p_list.add_run(desc)

    # Section 4: Theory Defense
    h = doc.add_heading(level=1)
    h.add_run("4. Lý Thuyết Cần Nắm Vững Khi Trả Lời Phản Biện").font.color.rgb = PRIMARY_COLOR
    
    theories = [
        ("Khái niệm mạng CNN và Depthwise Convolution:", " Mạng nơ-ron tích chập tự động trích xuất đặc trưng hình ảnh. Lớp Depthwise tách biệt việc tích chập không gian và tích chập kênh màu giúp mạng nhẹ hơn 9 lần."),
        ("Khái niệm Biến đổi Fourier:", " Là công cụ toán học chuyển ảnh từ miền không gian màu sắc sang miền tần số. Giúp phát hiện cấu trúc nhân tạo lặp đi lặp lại như lưới điểm ảnh iPad hay hạt mực giấy."),
        ("Cách tính độ chính xác và hàm Loss:", " Sử dụng hàm Cross-Entropy Loss để phạt AI nếu phân loại sai mặt thật/giả, kết hợp MSE Loss để ép AI học được cấu trúc phổ tần số Fourier giống ảnh mẫu.")
    ]
    for title, desc in theories:
        p_list = add_paragraph_with_spacing(doc, style='List Bullet')
        r_title = p_list.add_run(title)
        r_title.font.bold = True
        p_list.add_run(desc)

    doc.save("Phan_Cong_Thanh_Vien_1_AI_Engineer.docx")


def create_member_2_doc():
    doc = Document()
    sections = doc.sections
    for s in sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1.0)

    # Styles
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    PRIMARY_COLOR = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_COLOR = RGBColor(0x5B, 0x9B, 0xD5)

    # Header
    p_title = add_paragraph_with_spacing(doc, space_after=Pt(12))
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("BẢN PHÂN CÔNG NHIỆM VỤ THÀNH VIÊN 2: KỸ SƯ TỐI ƯU HÓA PHẦN CỨNG (HARDWARE OPTIMIZATION ENGINEER)")
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    p_role = add_paragraph_with_spacing(doc, space_after=Pt(18))
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role.add_run("Mức độ đóng góp: ").font.bold = True
    p_role.add_run("CHÍNH (LỚN NHẤT) - 40% khối lượng dự án").font.color.rgb = PRIMARY_COLOR

    # Section 1: Overview
    h = doc.add_heading(level=1)
    h.add_run("1. Mục Tiêu Chính").font.color.rgb = PRIMARY_COLOR
    p = add_paragraph_with_spacing(doc)
    p.add_run("Thành viên này chịu trách nhiệm biến 'bộ não' AI thô chạy chậm trên CPU thành một hệ thống chạy siêu tốc trên card đồ họa GPU NVIDIA. Nhiệm vụ chính xoay quanh lập trình tăng tốc luồng dữ liệu phần cứng, biên dịch mô hình sang dạng máy TensorRT, loại bỏ độ trễ của CPU và tối ưu hóa kích thước mô hình bằng lượng tử hóa chọn lọc.")

    # Section 2: Files
    h = doc.add_heading(level=1)
    h.add_run("2. Các Tệp Tin Đảm Nhận Trong Dự Án").font.color.rgb = PRIMARY_COLOR
    
    files_data = [
        ("src/inference/preprocess_cuda.py", "Lớp GPUPayloadPreprocessor thực hiện tiền xử lý (cắt ảnh, resize, chuẩn hóa) trực tiếp trên GPU VRAM."),
        ("src/inference/inference_trt.py", "Lớp TensorRTEngineWrapper nạp file engine, quản lý luồng CUDA stream và thực thi ghi/chạy kịch bản CUDA Graphs."),
        ("scripts/export_tensorrt.py", "Lập trình quy trình biên dịch mô hình sang định dạng TensorRT FP16 và lượng tử hóa tĩnh chọn lọc sang INT8 QDQ.")
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Đường Dẫn Tệp Tin"
    hdr_cells[1].text = "Vai Trò Chi Tiết Trong Mã Nguồn"
    for cell in hdr_cells:
        set_cell_background(cell, "F2F5F8")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = PRIMARY_COLOR

    for path, desc in files_data:
        row_cells = table.add_row().cells
        row_cells[0].text = path
        row_cells[1].text = desc
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            if idx == 0:
                cell.paragraphs[0].runs[0].font.name = 'Courier New'
                cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                cell.paragraphs[0].runs[0].font.bold = True

    # Section 3: Tasks
    h = doc.add_heading(level=1)
    h.add_run("3. Nhiệm Vụ Chi Tiết Từng Bước").font.color.rgb = PRIMARY_COLOR
    
    tasks = [
        ("Đẩy Tiền xử lý lên GPU:", " Lập trình lớp GPUPayloadPreprocessor sử dụng tensor PyTorch trên CUDA để toàn bộ khâu cắt ảnh và resize diễn ra trên GPU, tránh truyền dữ liệu đi lại qua cổng PCIe."),
        ("Biên dịch mô hình TensorRT:", " Viết lệnh cấu hình TensorRT config để chuyển mô hình .onnx thành file chạy .engine tối ưu riêng cho GPU RTX 3050."),
        ("Triển khai CUDA Graphs:", " Thiết lập cơ chế ghi lại luồng xử lý ở khung hình đầu tiên và chạy lại kịch bản tĩnh đó ở các khung hình sau để CPU không phải tốn thời gian gọi GPU lẻ tẻ."),
        ("Thực hiện Lượng tử hóa chọn lọc (Selective Quantization):", " Lọc bỏ các lớp tích chập nhạy cảm (Depthwise Conv, SE Modules) để chạy ở FP16 nhằm bảo toàn độ chính xác 84.40%, và chỉ ép các lớp Conv 1x1 thông thường chạy ở INT8 để giảm khối lượng tính toán.")
    ]
    for title, desc in tasks:
        p_list = add_paragraph_with_spacing(doc, style='List Bullet')
        r_title = p_list.add_run(title)
        r_title.font.bold = True
        p_list.add_run(desc)

    # Section 4: Theory Defense
    h = doc.add_heading(level=1)
    h.add_run("4. Lý Thuyết Cần Nắm Vững Khi Trả Lời Phản Biện").font.color.rgb = PRIMARY_COLOR
    
    theories = [
        ("PCIe Bottleneck (Nghẽn cổng truyền tải dữ liệu):", " Giải thích rằng truyền ảnh liên tục từ RAM lên VRAM qua PCIe rất chậm. GPU Preprocessing giúp chuyển ảnh thô 1 lần duy nhất lên GPU, loại bỏ điểm nghẽn này."),
        ("CUDA Graphs & CPU Launch Overhead:", " Giải thích rằng CPU ra lệnh cho GPU chạy từng lớp sẽ rất lâu. CUDA Graphs ghi lại kịch bản để GPU tự chạy một mạch, giải phóng CPU hoàn toàn."),
        ("Lượng tử hóa chọn lọc (Selective Quantization):", " Giải thích tại sao không lượng tử hóa toàn bộ (vì lớp Depthwise có ít tham số nên rất nhạy cảm). Giữ lớp nhạy cảm ở FP16 và nén lớp tính toán nặng sang INT8 giúp bảo toàn độ chính xác 100%.")
    ]
    for title, desc in theories:
        p_list = add_paragraph_with_spacing(doc, style='List Bullet')
        r_title = p_list.add_run(title)
        r_title.font.bold = True
        p_list.add_run(desc)

    doc.save("Phan_Cong_Thanh_Vien_2_Hardware_Engineer.docx")


def create_member_3_doc():
    doc = Document()
    sections = doc.sections
    for s in sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1.0)

    # Styles
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    PRIMARY_COLOR = RGBColor(0x1F, 0x4E, 0x79)
    SECONDARY_COLOR = RGBColor(0x5B, 0x9B, 0xD5)

    # Header
    p_title = add_paragraph_with_spacing(doc, space_after=Pt(12))
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("BẢN PHÂN CÔNG NHIỆM VỤ THÀNH VIÊN 3: KỸ SƯ TÍCH HỢP & ĐẢM BẢO CHẤT LƯỢNG (INTEGRATION & QA ENGINEER)")
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    p_role = add_paragraph_with_spacing(doc, space_after=Pt(18))
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role.add_run("Mức độ đóng góp: ").font.bold = True
    p_role.add_run("TRUNG BÌNH (MEDIUM) - 20% khối lượng dự án").font.color.rgb = PRIMARY_COLOR

    # Section 1: Overview
    h = doc.add_heading(level=1)
    h.add_run("1. Mục Tiêu Chính").font.color.rgb = PRIMARY_COLOR
    p = add_paragraph_with_spacing(doc)
    p.add_run("Thành viên này đóng vai trò cầu nối, tích hợp 'bộ não' AI của Thành viên 1 và 'động cơ tối ưu' của Thành viên 2 thành một ứng dụng hoàn chỉnh có thể chạy thực tế. Nhiệm vụ chính xoay quanh xây dựng giao diện Webcam thời gian thực, lập trình các chương trình đo đạc hiệu năng phần cứng, kiểm thử độ chính xác theo tiêu chuẩn quốc tế ISO và viết tài liệu hướng dẫn.")

    # Section 2: Files
    h = doc.add_heading(level=1)
    h.add_run("2. Các Tệp Tin Đảm Nhận Trong Dự Án").font.color.rgb = PRIMARY_COLOR
    
    files_data = [
        ("demo_trt.py", "Chương trình Demo chính xử lý luồng Webcam từ camera, nạp mô hình TensorRT, vẽ khung xanh/đỏ và hiển thị FPSOverlay."),
        ("scripts/evaluate_metrics.py", "Chương trình kiểm thử đánh giá các chỉ số độ chính xác học thuật của liveness (APCER, BPCER, ACER)."),
        ("scripts/benchmark_visual.py", "Đo đạc chi tiết thời gian trễ của CPU/GPU/TensorRT và tự động sinh biểu đồ so sánh hình ảnh."),
        ("benchmark_hardware.py", "Giám sát tài nguyên phần cứng, ghi lại tải CPU, GPU và VRAM trong quá trình chạy."),
        ("scripts/prepare_calib_data.py", "Trích xuất và chuẩn bị 250 ảnh mặt khuôn mặt từ tập dữ liệu gốc để làm tập dữ liệu kiểm thử liveness.")
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Đường Dẫn Tệp Tin"
    hdr_cells[1].text = "Vai Trò Chi Tiết Trong Mã Nguồn"
    for cell in hdr_cells:
        set_cell_background(cell, "F2F5F8")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = PRIMARY_COLOR

    for path, desc in files_data:
        row_cells = table.add_row().cells
        row_cells[0].text = path
        row_cells[1].text = desc
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            if idx == 0:
                cell.paragraphs[0].runs[0].font.name = 'Courier New'
                cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                cell.paragraphs[0].runs[0].font.bold = True

    # Section 3: Tasks
    h = doc.add_heading(level=1)
    h.add_run("3. Nhiệm Vụ Chi Tiết Từng Bước").font.color.rgb = PRIMARY_COLOR
    
    tasks = [
        ("Tích hợp mô hình phát hiện mặt YuNet:", " Sử dụng thư viện OpenCV để phát hiện mặt người từ khung hình camera thô, lấy tọa độ bounding box để chuyển sang mô hình liveness phân loại."),
        ("Xây dựng chương trình Demo thời gian thực (demo_trt.py):", " Lập trình vòng lặp camera 30 FPS, hiển thị nhãn xanh lá REAL / đỏ SPOOF quanh khuôn mặt, tạo bảng overlay thông số thiết bị biên."),
        ("Lập trình script đo lường độ chính xác (evaluate_metrics.py):", " Thiết lập quy trình tính toán các chỉ số lỗi APCER (lọt ảnh giả mạo) và BPCER (chặn nhầm người thật) để báo cáo."),
        ("Lập trình script trực quan hóa biểu đồ (benchmark_visual.py):", " Sử dụng matplotlib để tự động vẽ biểu đồ cột so sánh FPS và Latency, lưu lại dưới dạng file hình ảnh để dán vào báo cáo của nhóm."),
        ("Khắc phục các lỗi vận hành phần mềm:", " Giải quyết lỗi không mở được cửa sổ camera trên Windows và tự động nạp trước các file DLL của CUDA.")
    ]
    for title, desc in tasks:
        p_list = add_paragraph_with_spacing(doc, style='List Bullet')
        r_title = p_list.add_run(title)
        r_title.font.bold = True
        p_list.add_run(desc)

    # Section 4: Theory Defense
    h = doc.add_heading(level=1)
    h.add_run("4. Lý Thuyết Cần Nắm Vững Khi Trả Lời Phản Biện").font.color.rgb = PRIMARY_COLOR
    
    theories = [
        ("Tiêu chuẩn đánh giá ISO/IEC 30107-3:", " Hiểu rõ ý nghĩa của APCER (Tỉ lệ lọt ảnh giả mạo), BPCER (Tỉ lệ chặn nhầm người thật) và ACER (Tỉ lệ lỗi trung bình)."),
        ("Cách hệ thống camera và YuNet hoạt động:", " Giải thích cách YuNet tìm khuôn mặt nhanh chóng trên CPU bằng mạng CNN siêu nhỏ, làm tiền đề để cắt mặt đưa vào phân loại liveness."),
        ("Cách đọc biểu đồ so sánh hiệu năng:", " Giải thích được tại sao TRT FP16 đạt 250 FPS là cấu hình mượt nhất cho camera nhờ loại bỏ độ trễ của CPU.")
    ]
    for title, desc in theories:
        p_list = add_paragraph_with_spacing(doc, style='List Bullet')
        r_title = p_list.add_run(title)
        r_title.font.bold = True
        p_list.add_run(desc)

    doc.save("Phan_Cong_Thanh_Vien_3_QA_Engineer.docx")


if __name__ == "__main__":
    create_member_1_doc()
    create_member_2_doc()
    create_member_3_doc()
    print("[OK] All 3 roles Word documents created successfully.")
